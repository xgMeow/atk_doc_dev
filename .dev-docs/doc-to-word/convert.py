"""
VuePress 文档转 Word (.docx) 工具
用法: python convert.py            # 转换全部文档
      python convert.py 02-案例教程  # 只转换指定目录
"""
import sys
import re
import shutil
import subprocess
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# 路径：工具目录 / doc-to-word
TOOL_DIR = Path(__file__).parent.resolve()
# 仓库根目录（doc-to-word 的上级）
REPO_ROOT = TOOL_DIR.parent.resolve()
# 输出到工具目录下的 word_output
OUTPUT_DIR = TOOL_DIR / "word_output"
# 参考模板
REFERENCE_DOC = TOOL_DIR / "reference.docx"


def find_pandoc():
    """查找 pandoc 可执行文件"""
    # 优先从 PATH 中查找
    path = shutil.which("pandoc")
    if path:
        return path
    # Windows 默认安装路径
    for p in [
        "C:/Program Files/Pandoc/pandoc.exe",
        "C:/Program Files (x86)/Pandoc/pandoc.exe",
    ]:
        if Path(p).exists():
            return p
    return None


# ── 预处理 ───────────────────────────────────────────

def strip_vue_blocks(text: str) -> str:
    text = re.sub(r'<script[^>]*setup[^>]*>.*?</script>', '', text, flags=re.DOTALL)
    return text


def strip_vue_components(text: str) -> str:
    text = re.sub(r'<[A-Z][A-Za-z]*\b[^>]*?/\s*>', '', text)
    return text


def handle_containers(text: str) -> str:
    text = re.sub(
        r'^:::\s*(warning|tip|danger|details|info|note|important|caution)\s*(.*?)\s*$',
        r'> **\1**: \2\n>', text, flags=re.MULTILINE,
    )
    text = re.sub(r'^:::$', '', text, flags=re.MULTILINE)
    return text


def fix_kbd_tags(text: str) -> str:
    text = re.sub(r'<kbd>(.+?)</kbd>', r'`\1`', text, flags=re.DOTALL)
    return text


def resolve_include(text: str, base_dir: Path) -> str:
    pattern = re.compile(r'<!--\s*@include:\s*(.+?)\s*-->')
    while True:
        m = pattern.search(text)
        if not m:
            break
        rel_path = m.group(1).strip()
        inc_file = (base_dir / rel_path).resolve()
        if inc_file.exists():
            inc_content = inc_file.read_text(encoding="utf-8")
            text = text[:m.start()] + "\n" + inc_content + "\n" + text[m.end():]
        else:
            text = text[:m.start()] + text[m.end():]
    return text


def fix_image_paths(text: str, md_dir: Path) -> str:
    def resolve_path(raw_src: str) -> str:
        src_path = Path(raw_src)
        if src_path.is_absolute():
            return raw_src
        resolved = (md_dir / src_path).resolve()
        if resolved.exists():
            return str(resolved)
        for f in REPO_ROOT.rglob(src_path.name):
            return str(f.resolve())
        return raw_src

    text = re.sub(
        r'!\[([^\]]*)\]\(([^)]+)\)',
        lambda m: f'![{m.group(1)}]({resolve_path(m.group(2))})',
        text,
    )
    text = re.sub(
        r'(<img[^>]*src=")([^"]+)(")',
        lambda m: f'{m.group(1)}{resolve_path(m.group(2))}{m.group(3)}',
        text,
    )
    return text


def strip_yaml_frontmatter(text: str) -> str:
    text = text.strip()
    if text.startswith('---'):
        second = text.find('---', 3)
        if second != -1:
            text = text[second + 3:].strip()
    return text


def preprocess_markdown(md_path: Path) -> str:
    text = md_path.read_text(encoding="utf-8")
    text = strip_yaml_frontmatter(text)
    text = resolve_include(text, md_path.parent)
    text = strip_vue_blocks(text)
    text = strip_vue_components(text)
    text = handle_containers(text)
    text = fix_kbd_tags(text)
    text = fix_image_paths(text, md_path.parent)
    return text


# ── 后处理 ───────────────────────────────────────────

def postprocess_docx(docx_path: Path):
    """修正格式: 标题仅加粗+大小, 表格居中, 去除斜体"""
    doc = Document(str(docx_path))

    for p in doc.paragraphs:
        style_name = (p.style.name if p.style else "")
        for run in p.runs:
            run.font.italic = False
            rpr = run._element.get_or_add_rPr()
            rFonts = rpr.find(qn("w:rFonts"))
            if rFonts is None:
                rFonts = rpr.makeelement(qn("w:rFonts"), {})
                rpr.insert(0, rFonts)
            if not rFonts.get(qn("w:eastAsia")):
                rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        if style_name.startswith("Heading"):
            pPr = p._element.get_or_add_pPr()
            pBdr = pPr.find(qn("w:pBdr"))
            if pBdr is not None:
                pPr.remove(pBdr)

    for i in range(1, 7):
        try:
            hs = doc.styles[f"Heading {i}"]
            hs.font.italic = False
            hs.font.color.rgb = RGBColor(0, 0, 0)
            rpr = hs.element.get_or_add_rPr()
            rFonts = rpr.find(qn("w:rFonts"))
            if rFonts is None:
                rFonts = rpr.makeelement(qn("w:rFonts"), {})
                rpr.insert(0, rFonts)
            rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        except KeyError:
            pass

    for table in doc.tables:
        tbl = table._tbl
        tbl_pr = tbl.find(qn("w:tblPr"))
        if tbl_pr is None:
            tbl_pr = OxmlElement("w:tblPr")
            tbl.insert(0, tbl_pr)
        jc = tbl_pr.find(qn("w:jc"))
        if jc is None:
            jc = OxmlElement("w:jc")
            tbl_pr.append(jc)
        jc.set(qn("w:val"), "center")
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.italic = False

    doc.save(str(docx_path))


# ── 主流程 ───────────────────────────────────────────

def convert_one(md_path: Path, out_dir: Path, pandoc: str) -> bool:
    clean_md = preprocess_markdown(md_path)

    if len(clean_md.strip()) < 50:
        print(f"  [SKIP] {md_path.name}")
        return False

    rel_path = md_path.relative_to(REPO_ROOT)
    out_file = out_dir / rel_path.with_suffix(".docx")
    out_file.parent.mkdir(parents=True, exist_ok=True)

    tmp_md = out_file.with_suffix(".tmp.md")
    tmp_md.write_text(clean_md, encoding="utf-8")

    cmd = [pandoc, str(tmp_md), "-o", str(out_file),
           "--from=markdown", "--to=docx",
           "--metadata", "lang=zh-CN", "--standalone"]
    if REFERENCE_DOC.exists():
        cmd += ["--reference-doc", str(REFERENCE_DOC)]

    try:
        subprocess.run(cmd, check=True, capture_output=True, timeout=120)
    except subprocess.CalledProcessError as e:
        stderr = e.stderr.decode("utf-8", errors="replace")
        if "ERROR" in stderr:
            print(f"  [ERR] {rel_path}: {stderr[:200]}")
    finally:
        tmp_md.unlink(missing_ok=True)

    try:
        postprocess_docx(out_file)
    except Exception:
        pass

    print(f"  [OK] {rel_path}")
    return True


def main(target_dir: str = ""):
    pandoc = find_pandoc()
    if not pandoc:
        print("错误: 未找到 Pandoc，请先安装:")
        print("  winget install JohnMacFarlane.Pandoc")
        print("  或下载: https://pandoc.org/installing.html")
        sys.exit(1)
    print(f"Pandoc: {pandoc}")

    scan_dir = REPO_ROOT / target_dir if target_dir else REPO_ROOT
    if not scan_dir.exists():
        print(f"错误: 目录不存在 - {scan_dir}")
        sys.exit(1)

    # 先检查是否需要构建
    build_dir = REPO_ROOT / ".dist" / "standalone"
    if not build_dir.exists():
        print("未找到构建产物，正在构建...")
        subprocess.run(["npm", "run", "build:standalone"], cwd=str(REPO_ROOT), check=True)

    excluded = {"node_modules", ".git", ".vuepress", ".dist", ".claude", ".trae",
                ".idea", ".vscode", ".dev-docs", "doc-to-word", ".include"}
    files = []
    for f in scan_dir.rglob("*.md"):
        if not any(e in f.parts for e in excluded):
            files.append(f)

    root_skip = {"README.md", "COMMIT.md"}
    files = [f for f in files if f.name not in root_skip]

    if not files:
        print("未找到 Markdown 文件")
        return

    out_dir = OUTPUT_DIR
    out_dir.mkdir(parents=True, exist_ok=True)

    print(f"\nMD → DOCX: {len(files)} 个文件")
    print(f"输出目录: {out_dir}\n")

    ok = 0
    for fp in sorted(files):
        try:
            if convert_one(fp, out_dir, pandoc):
                ok += 1
        except Exception as e:
            print(f"  [ERR] {fp.relative_to(REPO_ROOT)}: {e}")

    print(f"\n完成! {ok}/{len(files)}")
    print(f"输出: {out_dir}")


if __name__ == "__main__":
    if len(sys.argv) > 1:
        main(sys.argv[1])
    else:
        main()
